import io
import re
from datetime import datetime

import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment

# Excel/XML'de geçersiz kontrol karakterleri (openpyxl IllegalCharacterError önler)
_ILLEGAL_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")

def _cell(val):
    """Hücre değerini Excel-güvenli stringe dönüştürür."""
    if val is None:
        return ""
    if isinstance(val, str):
        return _ILLEGAL_RE.sub("", val)
    return val

from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer

from docx import Document
from docx.shared import Pt, RGBColor

_MARKA_RENK   = "1A7A4E"   # koyu yeşil — Excel/Word başlık
_MARKA_HEX    = "#1A7A4E"
_SATIR_RENK   = "F0FAF4"   # açık yeşil — alternatif satır


def _bugun() -> str:
    return datetime.now().strftime("%d.%m.%Y %H:%M")


# ── Excel ──────────────────────────────────────────────────────────────────

def veri_to_excel(baslik: str, veri: list[list], sutunlar: list[str]) -> bytes:
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = baslik[:31]

    # Üst bilgi
    ws.append([baslik])
    ws.append([f"Oluşturma tarihi: {_bugun()}"])
    ws.append([])

    # Sütun başlıkları (4. satır)
    ws.append(sutunlar)
    for cell in ws[4]:
        cell.font      = Font(bold=True, color="FFFFFF", size=10)
        cell.fill      = PatternFill("solid", fgColor=_MARKA_RENK)
        cell.alignment = Alignment(horizontal="center", vertical="center")

    # Başlık satırı büyük font
    ws["A1"].font = Font(bold=True, size=14, color=_MARKA_RENK)

    # Veri satırları
    for idx, satir in enumerate(veri, start=5):
        ws.append([_cell(v) for v in satir])
        if idx % 2 == 0:
            for cell in ws[idx]:
                cell.fill = PatternFill("solid", fgColor=_SATIR_RENK)

    # Sütun genişlikleri otomatik
    for col in ws.columns:
        max_len = max((len(str(c.value or "")) for c in col), default=10)
        ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 50)

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ── PDF ────────────────────────────────────────────────────────────────────

def veri_to_pdf(baslik: str, veri: list[list], sutunlar: list[str]) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=40, bottomMargin=40)
    styles = getSampleStyleSheet()
    elems = []

    elems.append(Paragraph(baslik, styles["Title"]))
    elems.append(Paragraph(f"Oluşturma tarihi: {_bugun()}", styles["Normal"]))
    elems.append(Spacer(1, 12))

    tablo_veri = [sutunlar] + [[str(h) for h in row] for row in veri]
    tablo = Table(tablo_veri, repeatRows=1)
    tablo.setStyle(TableStyle([
        ("BACKGROUND",   (0, 0), (-1, 0),  colors.HexColor(_MARKA_HEX)),
        ("TEXTCOLOR",    (0, 0), (-1, 0),  colors.white),
        ("FONTNAME",     (0, 0), (-1, 0),  "Helvetica-Bold"),
        ("FONTSIZE",     (0, 0), (-1, -1), 8),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F0FAF4")]),
        ("GRID",         (0, 0), (-1, -1), 0.4, colors.HexColor("#CCCCCC")),
        ("ALIGN",        (0, 0), (-1, -1), "CENTER"),
        ("VALIGN",       (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING",   (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING",(0, 0), (-1, -1), 5),
    ]))
    elems.append(tablo)
    doc.build(elems)
    return buf.getvalue()


# ── Word ───────────────────────────────────────────────────────────────────

def veri_to_word(baslik: str, icerik: str, veri: list[list], sutunlar: list[str]) -> bytes:
    doc = Document()

    baslik_para = doc.add_heading(baslik, 0)
    if baslik_para.runs:
        baslik_para.runs[0].font.color.rgb = RGBColor(26, 122, 78)

    doc.add_paragraph(f"Oluşturma tarihi: {_bugun()}")
    if icerik:
        doc.add_paragraph(icerik)

    if veri and sutunlar:
        tablo = doc.add_table(rows=1, cols=len(sutunlar))
        tablo.style = "Table Grid"
        hdr = tablo.rows[0].cells
        for i, baslik_col in enumerate(sutunlar):
            hdr[i].text = baslik_col
            run = hdr[i].paragraphs[0].runs[0]
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

        for satir in veri:
            row = tablo.add_row().cells
            for i, deger in enumerate(satir):
                row[i].text = str(deger)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
